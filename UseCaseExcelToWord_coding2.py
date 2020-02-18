# -*- coding: utf-8 -*- 
"""
@File   :UseCaseExcelToWord_coding.py
@Software: PyCharm
@Author :张鹏
@Email  :1050035126@qq.com
@Date   :2019/11/25/0025 16:07
@Version:1.0
@Desc   :

解析coding的测试用例导出excel的测试用例到word中
python环境:
  python3
需要的库：pip install **

 1.python-docx   (生成word)
 2.csv           (读取用例导出excel)

需要预定义用例标识符的英文名称

"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import csv


def cleanStrList(strList):
    """
    整理字符串的内容
    :param strList:
    :return:
    """
    result = []
    for item in strList:
        if not item:
            continue
        item = re.sub(' +', ' ', item)
        item = re.sub('\\n$', '', item)
        item = re.sub("[\n\t]", "，", item)
        item = item.replace("，，", "，")
        item = item.replace("，”", "”")
        result.append(item)

    return result


def processStepToWord(table, stepListStr, stepResultListStr, rowNum, colNum):
    """
    解析用例中的每一步骤到步骤描述table中
    :param table:
    :param stepListStr:
    :param stepResultListStr:
    :return:
    """
    # 设置表头
    cells0 = table.rows[0].cells
    cells0[0].paragraphs[0].add_run("步骤").bold = True
    cells0[1].paragraphs[0].add_run("操作描述").bold = True
    cells0[2].paragraphs[0].add_run("预期结果").bold = True

    stepList = re.split('\d+\.', stepListStr)
    stepResultList = re.split('\d+\.', stepResultListStr)

    stepList = cleanStrList(stepList)
    stepResultList = cleanStrList(stepResultList)

    assert len(stepList) == len(stepResultList)

    if len(stepList) != len(stepResultList):
        return 0

    if len(stepList) == 0 or len(stepResultList) == 0:
        return 0

    for item in stepList:
        if item == "检查所有字段，1":
            print(1)
        assert item != ""
    for item in stepResultList:
        assert item != ""

    for i in range(len(stepList)):
        row = table.add_row()
        cells = row.cells
        # tempRowIndex = tempRow._index  # 当前行的序号
        cells[0].text = str(i + 1)
        cells[1].text = stepList[i]
        cells[2].text = stepResultList[i]

    # 设置table 中文字格式
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = u'宋体'
                    font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    col_width_dic = {0: 1,
                     1: 4.5,
                     2: 5.5
                     }
    for row_num in range(rowNum):
        for col_num in range(colNum):
            table.cell(row_num, col_num).width = Inches(col_width_dic[col_num])


def confirmUniqueSignal(signal):
    """
    确定标识符的唯一性
    :param signal:
    :return:
    """
    global signalStrList
    signalArray = [signal, 0]

    while True:
        if signalArray in signalStrList:
            signalArray = [signal, signalArray[1] + 1]
        else:
            signalStrList.append(signalArray)
            break

    if signalArray[1] == 0:
        # print(signal)
        return signal
    else:
        # print(signal + str(signalArray[1]))
        return signal + str(signalArray[1])


def createUserCaseTableWord(doc, rows, cols, title, caseNum, preCondition):
    """
    创建用例table的表头
    :param doc:
    :param rows:
    :param cols:
    :param title:
    :param caseNum:
    :param preCondition:
    :return:
    """
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')

    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中

    # header
    # cells = table.rows[0].cells
    # cells[0].paragraphs[0].add_run("模块名").bold = True
    # cells[1].merge(cells[cols - 1]).text = "触发信息管理模块"

    cells = table.rows[0].cells
    cells[0].paragraphs[0].add_run("涉及的需求")
    cells[1].paragraphs[0].add_run("内部接口")

    cells = table.rows[1].cells
    cells[0].paragraphs[0].add_run("先决条件")
    cells[1].paragraphs[0].add_run(preCondition)

    cells = table.rows[2].cells
    cells[0].paragraphs[0].add_run("功能描述")
    cells[1].paragraphs[0].add_run(title)

    return table


def docAddParagraph(doc, text, fontSize=12, fontName="宋体", fontBold=False, center=False):
    """
    doc中增加一段文字
    :param doc:
    :param text:
    :param fontSize:
    :param fontName:
    :param center:
    :return:
    """
    para = doc.add_paragraph()
    if center:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = para.add_run(text)
    r.font.name = fontName
    r.font.size = Pt(fontSize)
    r.font.bold = fontBold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fontName)


def getUseCaseByProjectName(useCaseList, projectName):
    """
    获取对应项目名称的所有用例
    :param useCaseList:
    :param projectName:
    :return:
    """
    result = []
    for useCase in useCaseList:
        caseProject = useCase[0]
        if caseProject == projectName:
            result.append(useCase)

    return result


def generateCaseTable(doc, useCaseList, projectName, caseIndex1, caseIndex2):
    """
    word中生成测试用例表格
    :param doc:
    :param useCaseList:用例完整导出
    :param caseIndex1: 用例说明序号1   4  表 4-2后台操作日志记录动作列表
    :param caseIndex2: 用例说明序号2   2  表 4-2后台操作日志记录动作列表
    :return:
    """
    rows = 3
    cols = 2

    doc.add_heading("测试说明", 1)
    docAddParagraph(doc,
                    "本软件分为所外、所内两个部分，所内部分功能较多，其菜单如下表所示，"
                    "其中个人资料、年度提案、会员管理、会员组管理、后台管理员、角色管理、登录成功日志、"
                    "登录失败日志、站内信管理、编号、年度的功能，是本软件与提案征集与反馈软件共用的功能，"
                    "其测试见提案征集与反馈软件测试说明。剩下的菜单对应的功能测试见本文档，下面的测试用例按照菜单划分章节，"
                    "测试用例中的入口对应的就是菜单项。",
                    12, "宋体", fontBold=False, center=False)

    moduleNameList = []

    i = 0
    for useCase in useCaseList:
        i += 1

        # if i == 3:
        #     break

        processPercent = round((i / len(useCaseList)) * 100, 2)

        print("\r处理进度%s%%" % processPercent, end="")
        moduleIndexStr = str(i)
        if i < 10:
            moduleIndexStr = "0" + moduleIndexStr

        project = useCase[0]

        module = useCase[1]
        mouleNamePre = "%s-%s" % (signalPre, moduleSigDic[projectName][module])
        moduleName = mouleNamePre + " " + module

        if moduleName not in moduleNameList:
            doc.add_heading(moduleName, 2)
            moduleNameList.append(moduleName)

        caseNum = "1"
        title = useCase[4]
        preCondition = useCase[5]
        stepListStr = useCase[7]
        stepResultListStr = useCase[8]

        titleStr = "测试用例:%s %s-%s" % (title, mouleNamePre, moduleIndexStr)
        doc.add_heading(titleStr, 3)

        caseIndex2 += 1
        tableExplainText = '表 %s-%s %s测试用例' % (caseIndex1, caseIndex2, title)
        # print(tableExplainText)
        docAddParagraph(doc, tableExplainText, 10, "黑体", center=True)

        table = createUserCaseTableWord(doc, rows, cols, title, caseNum, preCondition)

        # 内部嵌入测试步骤表格
        stepRow = table.add_row()
        stepCells = stepRow.cells

        stepCells[0].paragraphs[0].add_run("测试规程")

        stepTableRowNum = 1
        stepTableColNum = 3
        stepTable = stepCells[1].add_table(rows=stepTableRowNum, cols=stepTableColNum)
        stepTable.style = 'Table Grid'

        stepCells[1].paragraphs[0].paragraph_format.line_spacing = Pt(1)  # 固定值18磅

        failResult = processStepToWord(stepTable, stepListStr, stepResultListStr, stepTableRowNum, stepTableColNum)

        if failResult == 0:
            print("错误:" + title)

        lastRow = table.add_row()
        lastCells = lastRow.cells
        lastCells[0].paragraphs[0].add_run("假设和约束")
        lastCells[1].paragraphs[0].add_run("无")

        # 设置table 中文字格式
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = u'宋体'
                        font.size = Pt(10.5)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        col_width_dic = {0: 0.5,
                         1: 9.5,
                         }
        for row_num in range(rows):
            for col_num in range(cols):
                table.cell(row_num, col_num).width = Inches(col_width_dic[col_num])


def setDocBodyStyle(doc):
    """
    设置word全局的样式
    :param doc:
    :return:
    """
    styles = doc.styles

    style = styles['Normal']
    font = style.font
    font.size = Pt(12)
    # font.bold = True
    font.name = u'宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def generateWord(outWordDir, useCaseList):
    """
    创建word文档
    :param outWordDir:word文件输出文件夹
    :param useCaseList:
    :return:
    """
    if not os.path.exists(outWordDir):
        os.makedirs(outWordDir)

    # 测试用例说明序号
    # 表 4-2后台操作日志记录动作列表
    caseIndex1 = 4
    caseIndex2 = 0

    projectList = []
    for caseItem in useCaseList:
        projectName = caseItem[0]
        if projectName not in projectList:
            projectList.append(projectName)

    for projectName in projectList:
        print("正在生成:" + projectName)
        doc = Document()
        setDocBodyStyle(doc)

        projectCaseList = getUseCaseByProjectName(useCaseList, projectName)

        generateCaseTable(doc, projectCaseList, projectName, caseIndex1, caseIndex2)

        # doc.add_page_break()

        docOutPath = outWordDir + "/" + projectName + "_测试用例_生成.docx"

        doc.save(docOutPath)

        print("\n保存路径:" + docOutPath)


def getUseCaseList(useCasePath):
    """
    读取coding 用例数据excel
    :param useCasePath:
    :return:
    """
    result = []
    try:
        with open(useCasePath, encoding="utf8") as csvfile:
            csv_reader = csv.reader(csvfile)  # 使用csv.reader读取csvfile中的文件
            header = next(csv_reader)  # 读取第一行每一列的标题
            for row in csv_reader:  # 将csv 文件中的数据保存到birth_data中
                result.append(row)
    except Exception as e:
        print("coding用例导出excel读取失败：" + str(e))

    return result


def checkInputPath(excelPath, wordOutDir):
    """
    检查输入的文件路径是否存在
    :param excelPath:
    :param wordOutDir:
    :return:
    """
    if not os.path.exists(excelPath):
        print("用例导出excel文件不存在，请确认文件路径")
        return False

    try:
        if not os.path.exists(wordOutDir):
            os.makedirs(wordOutDir)
    except:
        print("测试用例文件夹创建失败")
        return False

    return True


if __name__ == '__main__':
    # TODO:需要预定义用例的前缀
    signalPre = "KY-04"
    # TODO:定义模块名称的英文缩写
    moduleSigDic = {
        "科学运行信息与业务监管软件": {
            "触发信息管理模块": "TRIGGER",
            "载荷监测与异常管理模块": "PAYLOAD",
            "数据传输模块": "FTP",
            "系统管理模块": "SYSTEM",
        },
        "载荷运行管理软件": {
            "更新申请提交模块": "APPLY-POST",
            "更新申请处理模块": "APPLY-PROCESS",
            "更新异常跟踪模块": "EXCEPTION-TRACK",
            "配置管理模块": "CONFIG",
            "系统管理模块": "SYSTEM",
        }

    }

    # TODO:读取从coding上导出的csv格式的测试用例
    excelPath = r"C:\Users\10500\Desktop\新建文件夹\用例导出-20200114205831.csv"
    # TODO:测试用例word生成文件夹
    wordOutDir = r"C:\Users\10500\Desktop\新建文件夹/"

    if checkInputPath(excelPath, wordOutDir):
        # 读取用例excel获取用例列表
        useCaseList = getUseCaseList(excelPath)
        # 分析用例,生成word测试文档
        generateWord(wordOutDir, useCaseList)
