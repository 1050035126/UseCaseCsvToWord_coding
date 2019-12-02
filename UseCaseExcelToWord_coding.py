# -*- coding: utf-8 -*- 
"""
@File   :UseCaseExcelToWord_coding.py
@Software: PyCharm
@Author :张鹏
@Email  :1050035126@qq.com
@Date   :2019/11/25/0025 16:07
@Version:1.0
@Desc   :

解析coding的测试用例导出excel的测试用例到word中
python环境:
  python3 (需要联网：存在翻译模块)
需要的库：pip install **

 1.python-docx   (生成word)
 2.csv           (读取用例导出excel)
 3.translate     (翻译中文到英文)
"""

import re
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from translate import Translator
import csv


def cleanStrList(strList):
    """
    整理字符串的内容
    :param strList:
    :return:
    """
    result = []
    for item in strList:
        item = re.sub(' +', ' ', item)
        item = re.sub('\\n$', '', item)
        item = re.sub("[\n\t]", "，", item)
        item = item.replace("，，", "，")
        item = item.replace("，”", "”")
        result.append(item)

    return result


def processStepToWord(table, stepListStr, stepResultListStr):
    """
    解析用例中的每一步骤到table中
    :param table:
    :param stepListStr:
    :param stepResultListStr:
    :return:
    """
    stepList = re.split('\d\.', stepListStr)[1:]
    stepResultList = re.split('\d\.', stepResultListStr)[1:]

    stepList = cleanStrList(stepList)
    stepResultList = cleanStrList(stepResultList)
    for i in range(len(stepList)):
        row = table.add_row()
        cells = row.cells
        # tempRowIndex = tempRow._index  # 当前行的序号
        cells[0].text = str(i + 1)
        cells[1].text = stepList[i]
        cells[2].text = stepResultList[i]

    # 设置table 中文字格式
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = u'宋体'
                    font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def confirmUniqueSignal(signal):
    """
    确定标识符的唯一性
    :param signal:
    :return:
    """
    global signalStrList
    signalArray = [signal, 0]

    while True:
        if signalArray in signalStrList:
            signalArray = [signal, signalArray[1] + 1]
        else:
            signalStrList.append(signalArray)
            break

    if signalArray[1] == 0:
        print(signal)
        return signal
    else:
        print(signal + str(signalArray[1]))
        return signal + str(signalArray[1])


def getSigalByTitle(title):
    """
    根据标题翻译获取用例的标识符

    阈值参数测试 KY-04-THR
    :param title:
    :return:
    """
    title = title.replace("-", " ")
    englishListStr = translator.translate(title)

    englishList = englishListStr.split(" ")
    result = "KY-04-"
    for item in englishList:
        result += item[0].upper()

    return confirmUniqueSignal(result)


def createUserCaseTableWord(doc, rows, cols, title, caseNum, preCondition):
    """
    创建用例table的表头
    :param doc:
    :param rows:
    :param cols:
    :param title:
    :param caseNum:
    :param preCondition:
    :return:
    """
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')

    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中

    # header
    # cells = table.rows[0].cells
    # cells[0].paragraphs[0].add_run("模块名").bold = True
    # cells[1].merge(cells[cols - 1]).text = "触发信息管理模块"

    cells = table.rows[0].cells
    cells[0].paragraphs[0].add_run("功能名").bold = True
    cells[1].merge(cells[cols - 1]).text = title

    cells = table.rows[1].cells
    cells[0].paragraphs[0].add_run("标识符").bold = True
    cells[1].merge(cells[cols - 1]).text = getSigalByTitle(title)

    cells = table.rows[2].cells
    cells[0].paragraphs[0].add_run("前置条件").bold = True
    cells[1].merge(cells[cols - 1]).text = preCondition

    cells = table.rows[3].cells
    cells[0].paragraphs[0].add_run("测试类型").bold = True
    cells[1].text = "功能测试"

    cells[2].paragraphs[0].add_run("测试工具").bold = True
    cells[3].merge(cells[cols - 1]).text = "无"

    cells = table.rows[4].cells
    cells[0].paragraphs[0].add_run("操作步骤").bold = True
    cells[1].paragraphs[0].add_run("操作描述").bold = True
    # cells[2].text = "数据"
    cells[2].paragraphs[0].add_run("预期结果").bold = True
    cells[3].paragraphs[0].add_run("实际结果").bold = True
    cells[4].paragraphs[0].add_run("测试状态").bold = True

    return table


def docAddParagraph(doc, text, fontSize=12, fontName="宋体", fontBold=False, center=False):
    """
    doc中增加一段文字
    :param doc:
    :param text:
    :param fontSize:
    :param fontName:
    :param center:
    :return:
    """
    para = doc.add_paragraph()
    if center:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = para.add_run(text)
    r.font.name = fontName
    r.font.size = Pt(fontSize)
    r.font.bold = fontBold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fontName)


def getUseCaseByProjectName(useCaseList, projectName):
    """
    获取对应项目名称的所有用例
    :param useCaseList:
    :param projectName:
    :return:
    """
    result = []
    for useCase in useCaseList:
        caseProject = useCase[0]
        if caseProject == projectName:
            result.append(useCase)

    return result


def generateCaseTable(doc, useCaseList, projectName, caseIndex1, caseIndex2):
    """
    word中生成测试用例表格
    :param doc:
    :param useCaseList:用例完整导出
    :param caseIndex1: 用例说明序号1   4  表 4-2后台操作日志记录动作列表
    :param caseIndex2: 用例说明序号2   2  表 4-2后台操作日志记录动作列表
    :return:
    """
    rows = 5
    cols = 5

    moduleList = []

    doc.add_heading("测试说明", 1)
    docAddParagraph(doc,
                    "本软件分为所外、所内两个部分，所内部分功能较多，其菜单如下表所示，"
                    "其中个人资料、年度提案、会员管理、会员组管理、后台管理员、角色管理、登录成功日志、"
                    "登录失败日志、站内信管理、编号、年度的功能，是本软件与提案征集与反馈软件共用的功能，"
                    "其测试见提案征集与反馈软件测试说明。剩下的菜单对应的功能测试见本文档，下面的测试用例按照菜单划分章节，"
                    "测试用例中的入口对应的就是菜单项。",
                    12, "宋体", fontBold=False, center=False)

    for useCase in useCaseList:
        project = useCase[0]

        module = useCase[1]
        # docAddParagraph(doc, '软件:' + project, 12, "黑体", fontBold=True, center=False)

        if module not in moduleList:
            moduleList.append(module)
            doc.add_heading(module, 2)
            # docAddParagraph(doc, "模块:" + module, 12, "黑体", fontBold=True, center=False)

        caseNum = "1"
        title = useCase[4]
        preCondition = useCase[5]
        stepListStr = useCase[7]
        stepResultListStr = useCase[8]

        # docAddParagraph(doc, '测试用例:' + title, 12, "黑体", fontBold=True, center=False)
        doc.add_heading('测试用例:' + title, 3)

        caseIndex2 += 1
        tableExplainText = '表 %s-%s %s测试用例' % (caseIndex1, caseIndex2, title)
        print(tableExplainText)
        docAddParagraph(doc, tableExplainText, 10, "黑体", center=True)

        table = createUserCaseTableWord(doc, rows, cols, title, caseNum, preCondition)
        processStepToWord(table, stepListStr, stepResultListStr)

        col_width_dic = {0: 0.5,
                         1: 2.5,
                         2: 4,
                         3: 0.5,
                         4: 0.5}
        for col_num in range(5):
            table.cell(5, col_num).width = Inches(col_width_dic[col_num])


def setDocBodyStyle(doc):
    """
    设置word全局的样式
    :param doc:
    :return:
    """
    styles = doc.styles

    style = styles['Normal']
    font = style.font
    font.size = Pt(12)
    # font.bold = True
    font.name = u'宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def generateWord(outWordDir, useCaseList):
    """
    创建word文档
    :param outWordDir:word文件输出文件夹
    :param useCaseList:
    :return:
    """
    if not os.path.exists(outDocDir):
        os.makedirs(outDocDir)

    # 测试用例说明序号
    # 表 4-2后台操作日志记录动作列表
    caseIndex1 = 4
    caseIndex2 = 1

    projectList = []
    for caseItem in useCaseList:
        projectName = caseItem[0]
        if projectName not in projectList:
            projectList.append(projectName)

    for projectName in projectList:
        print("正在生成:" + projectName)
        doc = Document()
        setDocBodyStyle(doc)

        projectCaseList = getUseCaseByProjectName(useCaseList, projectName)

        # addTableSysMenuTable(doc, projectCaseList, projectName)

        generateCaseTable(doc, projectCaseList, projectName, caseIndex1, caseIndex2)

        # doc.add_page_break()

        docOutPath = outWordDir + "/" + projectName + "_测试用例_生成.docx"

        doc.save(docOutPath)

        print("保存路径:" + docOutPath)


def getUseCaseList(useCasePath):
    """
    读取coding 用例数据excel
    :param useCasePath:
    :return:
    """
    result = []
    with open(useCasePath, encoding="utf8") as csvfile:
        csv_reader = csv.reader(csvfile)  # 使用csv.reader读取csvfile中的文件
        birth_header = next(csv_reader)  # 读取第一行每一列的标题
        for row in csv_reader:  # 将csv 文件中的数据保存到birth_data中
            result.append(row)

    return result


if __name__ == '__main__':
    # 标识符	KY-04-THR 列表 防止重复
    signalStrList = []
    # 定义汉翻英的翻译器 标识符生成
    translator = Translator(from_lang="chinese", to_lang="english")

    # 读取从coding上导出的csv格式的测试用例
    excelPath = r"C:\Users\10500\Desktop\完整用例导出.csv"
    # 测试用例word生成文件夹
    outDocDir = r"result2/"
    # 读取用例excel获取用例列表
    useCaseList = getUseCaseList(excelPath)
    # 分析用例,生成word测试文档
    generateWord(outDocDir, useCaseList)
